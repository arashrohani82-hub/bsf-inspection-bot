"""Render stable figure labels assigned before multipart splitting."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import inspection_bot as bot
from photo_numbering_runtime import _alpha_suffix


def _photo_label(photo: dict, fallback_group: int, fallback_index: int) -> str:
    figure = photo.get("_figure_number", fallback_group)
    letter_index = photo.get("_figure_letter_index", fallback_index)
    return f"{figure}{_alpha_suffix(int(letter_index))}"


def install_continuous_figure_runtime() -> None:
    original_insert = bot.insert_photo_groups

    def insert_photo_groups(doc, anchor_elem, groups, lang):
        has_stable_numbers = any(
            "_figure_number" in photo
            for group in groups
            for photo in group.get("photos", [])
        )
        if not has_stable_numbers:
            return original_insert(doc, anchor_elem, groups, lang)

        caption_k = "caption_fr" if lang == "fr" else "caption_en"
        insert_after = anchor_elem

        for fallback_group_num, group in enumerate(groups, 1):
            photos = group.get("photos", [])
            if not photos:
                continue
            caption = group.get(caption_k, "")
            labels = [
                _photo_label(photo, fallback_group_num, index)
                for index, photo in enumerate(photos)
            ]
            fig_label = (
                f"Fig. {labels[0]}"
                if len(labels) == 1
                else f"Fig. {labels[0]} à {labels[-1]}"
            )

            photos_padded = photos + [None] if len(photos) % 2 else photos
            pairs = [
                (photos_padded[i], photos_padded[i + 1])
                for i in range(0, len(photos_padded), 2)
            ]

            for pair_idx, (left, right) in enumerate(pairs):
                left_idx = pair_idx * 2
                right_idx = left_idx + 1
                left_label = labels[left_idx] if left_idx < len(labels) else ""
                right_label = labels[right_idx] if right_idx < len(labels) else ""

                tbl = doc.add_table(rows=1, cols=2)
                tbl.autofit = False
                tbl_pr = tbl._tbl.tblPr
                if tbl_pr is None:
                    tbl_pr = OxmlElement("w:tblPr")
                    tbl._tbl.insert(0, tbl_pr)

                tbl_jc = OxmlElement("w:jc")
                tbl_jc.set(qn("w:val"), "center")
                tbl_pr.append(tbl_jc)

                tbl_indent = OxmlElement("w:tblInd")
                tbl_indent.set(qn("w:w"), "0")
                tbl_indent.set(qn("w:type"), "dxa")
                tbl_pr.append(tbl_indent)

                tbl_width = OxmlElement("w:tblW")
                tbl_width.set(qn("w:w"), "8640")
                tbl_width.set(qn("w:type"), "dxa")
                tbl_pr.append(tbl_width)

                tbl_layout = OxmlElement("w:tblLayout")
                tbl_layout.set(qn("w:type"), "fixed")
                tbl_pr.append(tbl_layout)

                tbl_borders = OxmlElement("w:tblBorders")
                for border_name in (
                    "top", "left", "bottom", "right", "insideH", "insideV"
                ):
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "none")
                    tbl_borders.append(border)
                tbl_pr.append(tbl_borders)

                for cell in tbl.rows[0].cells:
                    cell.width = Inches(3.0)
                    cell.vertical_alignment = 1
                    tc_width = cell._tc.get_or_add_tcPr().first_child_found_in(
                        "w:tcW"
                    )
                    if tc_width is not None:
                        tc_width.set(qn("w:w"), "4320")
                        tc_width.set(qn("w:type"), "dxa")

                def fill_cell(cell, photo, label=""):
                    if photo is None:
                        return
                    img_path = photo.get("path")
                    img_para = cell.paragraphs[0]
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if not img_path or not Path(img_path).exists():
                        return

                    labelled_path = None
                    try:
                        with tempfile.NamedTemporaryFile(
                            suffix=".jpg", delete=False
                        ) as tmp:
                            labelled_path = tmp.name
                        bot.add_label_to_image(
                            img_path,
                            label,
                            labelled_path,
                            display_width_px=900,
                        )
                        img_para.add_run().add_picture(
                            labelled_path,
                            width=Inches(2.7),
                        )
                    except Exception as exc:
                        bot.log.warning(
                            "Could not insert continuously numbered photo: %s",
                            exc,
                        )
                        img_para.add_run("[image error]")
                    finally:
                        if labelled_path:
                            try:
                                Path(labelled_path).unlink(missing_ok=True)
                            except Exception:
                                pass

                fill_cell(tbl.rows[0].cells[0], left, left_label)
                fill_cell(tbl.rows[0].cells[1], right, right_label)

                tbl_el = tbl._tbl
                doc._body._body.remove(tbl_el)
                insert_after.addnext(tbl_el)
                insert_after = tbl_el

            cap_elem = OxmlElement("w:p")
            insert_after.addnext(cap_elem)
            for paragraph in doc.paragraphs:
                if paragraph._element is cap_elem:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(f"{fig_label} – {caption}")
                    run.italic = True
                    run.font.size = Pt(9)
                    break
            insert_after = cap_elem

            spacer = OxmlElement("w:p")
            insert_after.addnext(spacer)
            insert_after = spacer

    bot.insert_photo_groups = insert_photo_groups
