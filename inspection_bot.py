"""
BSF Inspections – Telegram Report Bot
With project database and inspection type selection
"""

import os, json, logging, subprocess, io, tempfile
from datetime import datetime
from pathlib import Path
import anthropic
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    ConversationHandler, ContextTypes, filters
)
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import report_profiles

logging.basicConfig(format="%(asctime)s | %(levelname)s | %(message)s", level=logging.INFO)
log = logging.getLogger(__name__)

TELEGRAM_TOKEN   = os.environ["TELEGRAM_TOKEN"]
ANTHROPIC_KEY    = os.environ["ANTHROPIC_API_KEY"]
anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)

(
    STATE_MAIN_MENU,
    STATE_INSPECTION_TYPE,
    STATE_PROJECT_SELECT,
    STATE_PLANS,
    STATE_DAVIT_DETAIL,
    STATE_PHOTO,
    STATE_GROUP_OR_ADD,
    STATE_ELEMENT_TYPE,
    STATE_ELEMENT_ID,
    STATE_PROBLEM,
    STATE_ELEMENT_STATUS,
    STATE_GROUP_CAPTION_FR,
    STATE_GROUP_CAPTION_EN,
    # Admin states
    STATE_ADMIN_MENU,
    STATE_ADMIN_PROJECT_NAME,
    STATE_ADMIN_PROJECT_ADDRESS,
    STATE_ADMIN_PROJECT_PLANS,
    STATE_ADMIN_PROJECT_DAVIT,
    STATE_CERTIFICATE_DECISION,
) = range(19)

ELEMENT_TYPES = [
    ["Anchor", "Davit"],
    ["Cable", "Base / Socket"],
    ["Roof", "Other"],
]


def get_element_types_for_session(session):
    return report_profiles.element_types(session.get("inspection_type"))


SEVERITY_MAP = {
    "critical": "🔴 Critical",
    "major":    "🟠 Major",
    "moderate": "🟡 Moderate",
    "minor":    "🟢 Minor",
    "ok":       "✅ OK / Acceptable",
}

BASE_DIR      = Path("/app")
SESSIONS_DIR  = BASE_DIR / "sessions";  SESSIONS_DIR.mkdir(exist_ok=True)
PHOTOS_DIR    = BASE_DIR / "photos";    PHOTOS_DIR.mkdir(exist_ok=True)
REPORTS_DIR   = BASE_DIR / "reports";   REPORTS_DIR.mkdir(exist_ok=True)
TEMPLATE_PATH = BASE_DIR / "Template.docx"
DB_PATH       = BASE_DIR / "projects.json"
# Fallback: if not in /app, try same directory as this script
if not DB_PATH.exists():
    DB_PATH = Path(__file__).parent / "projects.json"


# ── Database helpers ───────────────────────────────────────────────────────
def load_db():
    # Try /app/projects.json first, then same dir as script
    paths = [BASE_DIR / "projects.json", Path(__file__).parent / "projects.json"]
    for p in paths:
        if p.exists():
            return json.loads(p.read_text())
    return {"inspection_types": [], "projects": []}

def save_db(db):
    # Always save to a writable location
    for p in [BASE_DIR / "projects.json", Path(__file__).parent / "projects.json"]:
        try:
            p.write_text(json.dumps(db, ensure_ascii=False, indent=2))
            return
        except:
            continue



def get_projects():
    return load_db().get("projects", [])

def get_inspection_types():
    return load_db().get("inspection_types", [])



# ── Auto caption by inspection type + element type ─────────────────────────
CAPTION_MAP = {
    ("Ancrage 1 an", "Anchor"):        "Inspection visuelle des ancrages",
    ("Ancrage 1 an", "Davit"):         "Inspection visuelle des bossoirs",
    ("Ancrage 1 an", "Cable"):         "Inspection visuelle des lignes de vie",
    ("Ancrage 1 an", "Base / Socket"): "Inspection visuelle des socles",
    ("Ancrage 5 ans", "Anchor"):       "Essai de traction des ancrages",
    ("Ancrage 5 ans", "Davit"):        "Essai de traction des bossoirs",
    ("Ancrage 5 ans", "Cable"):        "Inspection et essai des lignes de vie",
    ("Ancrage 5 ans", "Base / Socket"):"Essai de traction des socles",
}

# ── Session helpers ────────────────────────────────────────────────────────
def session_path(chat_id): return SESSIONS_DIR / f"{chat_id}.json"
def load_session(chat_id):
    p = session_path(chat_id)
    return json.loads(p.read_text()) if p.exists() else {}
def save_session(chat_id, data):
    session_path(chat_id).write_text(json.dumps(data, ensure_ascii=False, indent=2))
def clear_session(chat_id):
    p = session_path(chat_id)
    if p.exists(): p.unlink()


# ── Replace placeholders ───────────────────────────────────────────────────
def replace_in_paragraph(para, replacements):
    full_text = para.text
    new_text  = full_text
    for key, val in replacements.items():
        new_text = new_text.replace(key, val)
    if new_text == full_text:
        return
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]: run.text = ""
    else:
        para.add_run(new_text)

def apply_replacements(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)


# ── Remove section if empty ────────────────────────────────────────────────
def clear_section_if_empty(doc, placeholder_text, also_remove_headers=None):
    target_elem = None
    for para in doc.paragraphs:
        if placeholder_text in para.text:
            target_elem = para._element
            break
    if target_elem is None:
        return
    to_remove = [target_elem]
    nxt = target_elem.getnext()
    count = 0
    while nxt is not None and count < 5:
        tag = nxt.tag.split("}")[-1] if "}" in nxt.tag else nxt.tag
        if tag == "p":
            text = "".join(t.text or "" for t in nxt.iter() if t.text)
            if text.strip() == "":
                to_remove.append(nxt); nxt = nxt.getnext(); count += 1
            else: break
        else: break
    if also_remove_headers:
        prev = target_elem.getprevious()
        count = 0
        while prev is not None and count < 6:
            tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
            if tag == "p":
                text = "".join(t.text or "" for t in prev.iter() if t.text)
                if text.strip() == "" or any(h in text for h in also_remove_headers):
                    to_remove.append(prev); prev = prev.getprevious(); count += 1
                else: break
            else: break
    for elem in to_remove:
        try: elem.getparent().remove(elem)
        except: pass


# ── Add label to image ─────────────────────────────────────────────────────
def add_label_to_image(img_path, label, output_path, display_width_px=800):
    """Add a readable figure number to the top-left corner of a photo."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.open(img_path).convert("RGB")

    # Resize first so the label keeps a consistent, readable size in Word.
    ratio = display_width_px / img.width
    img = img.resize(
        (display_width_px, max(1, int(img.height * ratio))),
        Image.LANCZOS,
    )
    draw = ImageDraw.Draw(img)
    font_size = max(64, min(84, int(img.width * 0.085)))
    font = None
    font_candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    ]
    nix_font_root = Path("/nix/store")
    if nix_font_root.exists():
        font_candidates.extend(
            nix_font_root.glob(
                "*dejavu-fonts*/share/fonts/truetype/DejaVuSans-Bold.ttf"
            )
        )

    for font_path in font_candidates:
        if not font_path.exists():
            continue
        try:
            font = ImageFont.truetype(str(font_path), font_size)
            break
        except Exception:
            continue

    if font is None:
        try:
            # Recent Pillow versions support a scalable bundled fallback.
            font = ImageFont.load_default(size=font_size)
        except TypeError:
            font = ImageFont.load_default()

    pad = max(10, font_size // 6)
    bbox = draw.textbbox((0, 0), label, font=font)
    width, height = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.rectangle(
        [0, 0, width + pad * 2, height + pad * 2],
        fill="white",
        outline="black",
        width=max(2, font_size // 18),
    )
    draw.text((pad, pad), label, fill="black", font=font)

    img.save(output_path, format="JPEG", quality=92)
    return output_path


# ── Insert single image ────────────────────────────────────────────────────
def insert_single_image(doc, anchor_para, img_path, width_inches=4.5, caption=""):
    insert_after = anchor_para._element
    img_elem = OxmlElement("w:p")
    insert_after.addnext(img_elem)
    for p in doc.paragraphs:
        if p._element is img_elem:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path and Path(img_path).exists():
                try: p.add_run().add_picture(img_path, width=Inches(width_inches))
                except Exception as e: p.add_run(f"[image error: {e}]")
            break
    insert_after = img_elem
    cap_elem = OxmlElement("w:p")
    insert_after.addnext(cap_elem)
    for p in doc.paragraphs:
        if p._element is cap_elem:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption); r.italic = True; r.font.size = Pt(9)
            break


# ── Insert plans vertically ────────────────────────────────────────────────
def insert_photos_vertical(doc, anchor_elem, photos, lang, img_width=5.5):
    caption_k    = "caption_fr" if lang == "fr" else "caption_en"
    insert_after = anchor_elem
    fig_num      = 1
    for photo in photos:
        ai = photo.get("ai", {})
        caption  = ai.get(caption_k, f"Plan {fig_num}")
        img_path = photo.get("path")
        img_elem = OxmlElement("w:p")
        insert_after.addnext(img_elem)
        for p in doc.paragraphs:
            if p._element is img_elem:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if img_path and Path(img_path).exists():
                    try: p.add_run().add_picture(img_path, width=Inches(img_width))
                    except: p.add_run("[image error]")
                break
        insert_after = img_elem
        cap_elem = OxmlElement("w:p")
        insert_after.addnext(cap_elem)
        for p in doc.paragraphs:
            if p._element is cap_elem:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"Fig. {fig_num} – {caption}")
                r.italic = True; r.font.size = Pt(9)
                break
        insert_after = cap_elem
        spacer = OxmlElement("w:p")
        insert_after.addnext(spacer)
        insert_after = spacer
        fig_num += 1


# ── Insert photo groups ────────────────────────────────────────────────────
def insert_photo_groups(doc, anchor_elem, groups, lang):
    caption_k = "caption_fr" if lang == "fr" else "caption_en"
    insert_after = anchor_elem
    group_num = 1

    for group in groups:
        photos = group.get("photos", [])
        caption = group.get(caption_k, "")
        n = len(photos)
        letters = [chr(ord("a") + i) for i in range(n)]
        fig_label = (
            f"Fig. {group_num}"
            if n == 1
            else f"Fig. {group_num}a à {group_num}{letters[-1]}"
        )

        photos_padded = photos + [None] if n % 2 else photos
        pairs = [
            (photos_padded[i], photos_padded[i + 1])
            for i in range(0, len(photos_padded), 2)
        ]

        for pair_idx, (left, right) in enumerate(pairs):
            left_idx = pair_idx * 2
            right_idx = left_idx + 1
            left_label = (
                f"{group_num}{letters[left_idx]}"
                if left_idx < len(letters)
                else ""
            )
            right_label = (
                f"{group_num}{letters[right_idx]}"
                if right_idx < len(letters)
                else ""
            )

            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            tblPr = tbl._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl._tbl.insert(0, tblPr)

            # Explicit table centering is needed because Word otherwise aligns
            # an inserted table to the left edge of the text area.
            tbl_jc = OxmlElement("w:jc")
            tbl_jc.set(qn("w:val"), "center")
            tblPr.append(tbl_jc)

            tbl_indent = OxmlElement("w:tblInd")
            tbl_indent.set(qn("w:w"), "0")
            tbl_indent.set(qn("w:type"), "dxa")
            tblPr.append(tbl_indent)

            tbl_width = OxmlElement("w:tblW")
            tbl_width.set(qn("w:w"), "8640")
            tbl_width.set(qn("w:type"), "dxa")
            tblPr.append(tbl_width)

            tbl_layout = OxmlElement("w:tblLayout")
            tbl_layout.set(qn("w:type"), "fixed")
            tblPr.append(tbl_layout)

            tbl_borders = OxmlElement("w:tblBorders")
            for border_name in (
                "top", "left", "bottom", "right", "insideH", "insideV"
            ):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tbl_borders.append(border)
            tblPr.append(tbl_borders)

            for cell in tbl.rows[0].cells:
                cell.width = Inches(3.0)
                cell.vertical_alignment = 1
                tc_width = cell._tc.get_or_add_tcPr().first_child_found_in(
                    "w:tcW"
                )
                if tc_width is not None:
                    tc_width.set(qn("w:w"), "4320")
                    tc_width.set(qn("w:type"), "dxa")

            def fill_cell(cell, photo, label=""):
                if photo is None:
                    return
                img_path = photo.get("path")
                img_para = cell.paragraphs[0]
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not img_path or not Path(img_path).exists():
                    return

                labelled_path = None
                try:
                    with tempfile.NamedTemporaryFile(
                        suffix=".jpg", delete=False
                    ) as tmp:
                        labelled_path = tmp.name
                    add_label_to_image(
                        img_path,
                        label,
                        labelled_path,
                        display_width_px=900,
                    )
                    img_para.add_run().add_picture(
                        labelled_path,
                        width=Inches(2.7),
                    )
                except Exception as exc:
                    log.warning("Could not insert labelled photo: %s", exc)
                    img_para.add_run("[image error]")
                finally:
                    if labelled_path:
                        try:
                            Path(labelled_path).unlink(missing_ok=True)
                        except Exception:
                            pass

            fill_cell(tbl.rows[0].cells[0], left, left_label)
            fill_cell(tbl.rows[0].cells[1], right, right_label)

            tbl_el = tbl._tbl
            doc._body._body.remove(tbl_el)
            insert_after.addnext(tbl_el)
            insert_after = tbl_el

        cap_elem = OxmlElement("w:p")
        insert_after.addnext(cap_elem)
        for paragraph in doc.paragraphs:
            if paragraph._element is cap_elem:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(f"{fig_label} – {caption}")
                run.italic = True
                run.font.size = Pt(9)
                break
        insert_after = cap_elem

        spacer = OxmlElement("w:p")
        insert_after.addnext(spacer)
        insert_after = spacer
        group_num += 1


# ── Claude Vision ──────────────────────────────────────────────────────────
def analyse_photo(
    image_bytes,
    element_type,
    location,
    problem,
    inspection_type="",
):
    import base64
    b64 = base64.standard_b64encode(image_bytes).decode()
    domain_context = report_profiles.ai_context(inspection_type)
    prompt = f"""You are an engineering inspection report assistant.
Inspection context: {domain_context}
Observed category: element={element_type}, location={location}.
Inspector note: {problem or 'infer only the visible condition from the image'}.
Do not claim that hidden conditions, measurements, destructive tests, or load
tests were verified unless the inspector note explicitly says so.
Respond ONLY with valid JSON:
{{"caption_fr":"...","caption_en":"...","severity":"critical|major|moderate|minor|ok"}}"""
    for attempt in range(3):
        try:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-20250514", max_tokens=300, timeout=60,
                messages=[{"role":"user","content":[
                    {"type":"image","source":{"type":"base64","media_type":"image/jpeg","data":b64}},
                    {"type":"text","text":prompt}
                ]}])
            raw = response.content[0].text.strip().replace("```json","").replace("```","").strip()
            return json.loads(raw)
        except Exception as e:
            log.warning(f"API attempt {attempt+1} failed: {e}")
            if attempt == 2: raise
            import time; time.sleep(3)



# ── DOCX to PDF ────────────────────────────────────────────────────────────
def docx_to_pdf(docx_path, pdf_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.units import cm
    import re
    doc    = Document(docx_path)
    story  = []
    styles = getSampleStyleSheet()
    for para in doc.paragraphs:
        text = re.sub(r"[^\x00-\xFF]", "", para.text.strip())
        if not text:
            story.append(Spacer(1, 0.2*cm)); continue
        if "Heading 1" in para.style.name: s = styles["Heading1"]
        elif "Heading" in para.style.name: s = styles["Heading2"]
        else: s = styles["Normal"]
        story.append(Paragraph(text, s))
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img = Image.open(io.BytesIO(rel.target_part.blob)).convert("RGB")
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                    img.save(tmp.name, "JPEG", quality=85)
                    w = min(15*cm, img.width * 0.026 * cm)
                    story.append(RLImage(tmp.name, width=w, height=w*img.height/img.width))
                    story.append(Spacer(1, 0.3*cm))
            except Exception: pass
    SimpleDocTemplate(pdf_path, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm).build(story)

# ── Report builder ─────────────────────────────────────────────────────────
def build_report(session, lang):
    inspection_type = session.get("inspection_type", "")
    template = report_profiles.template_path(BASE_DIR, inspection_type)
    doc          = Document(template)
    project      = session.get("project_name", "—")
    address      = session.get("address", "—")
    date         = session.get("date", datetime.today().strftime("%Y-%m-%d"))
    groups       = session.get("groups", [])
    plans        = session.get("plans", [])
    davit_detail = session.get("davit_detail")

    replacements = {
        "{{Project_Name}}": project,
        "{{Address_of _project }}": address,
        "{{Address_of_project}}": address,
        "{{Date }}": date,
        "{{Date}}": date,
        "{{caption}}": "",
    }
    replacements.update(
        report_profiles.report_replacements(inspection_type, groups)
    )
    apply_replacements(doc, replacements)

    if plans:
        for para in doc.paragraphs:
            if "{{Plans}}" in para.text:
                for run in para.runs: run.text = ""
                insert_photos_vertical(doc, para._element, plans, lang, img_width=5.5)
                break
    else:
        clear_section_if_empty(doc, "{{Plans}}", also_remove_headers=["Plans disponibles","Documents et références"])

    if davit_detail:
        for para in doc.paragraphs:
            if "{{Detail_davit" in para.text:
                for run in para.runs: run.text = ""
                cap = "Fig. 2 : Détail de configuration des bossoirs" if lang=="fr" else "Fig. 2 : Davit configuration detail"
                insert_single_image(doc, para, davit_detail, width_inches=4.5, caption=cap)
                break
    else:
        clear_section_if_empty(doc, "{{Detail_davit", also_remove_headers=["Fig. 1 : les plans","Fig. 2 : Détail","Rapports antérieurs"])

    photos_para = None
    for para in doc.paragraphs:
        if "{{Photos" in para.text:
            photos_para = para; break
    if photos_para is None:
        photos_para = doc.add_paragraph()
    for run in photos_para.runs: run.text = ""
    if groups:
        # Separate acceptable from non-acceptable photos
        acceptable_groups   = []
        intervention_groups = []

        for group in groups:
            acc_photos   = [p for p in group["photos"] if "Acceptable" in p.get("status","✅ Acceptable")]
            inter_photos = [p for p in group["photos"] if "Acceptable" not in p.get("status","✅ Acceptable")]

            if acc_photos:
                g = dict(group); g["photos"] = acc_photos
                acceptable_groups.append(g)
            if inter_photos:
                g = dict(group); g["photos"] = inter_photos
                intervention_groups.append(g)

        # Insert acceptable photos
        if acceptable_groups:
            insert_photo_groups(doc, photos_para._element, acceptable_groups, lang)

        # Insert intervention section
        if intervention_groups:
            # Add section title
            title_fr = "Éléments nécessitant une intervention"
            title_en = "Elements requiring intervention"
            title_text = title_fr if lang == "fr" else title_en

            # Find last inserted element to add after
            title_elem = OxmlElement("w:p")
            photos_para._element.addnext(title_elem)
            for p in doc.paragraphs:
                if p._element is title_elem:
                    r = p.add_run(title_text)
                    r.bold = True
                    r.font.size = Pt(11)
                    break

            insert_photo_groups(doc, title_elem, intervention_groups, lang)

    suffix = "FR" if lang=="fr" else "EN"
    profile_name = report_profiles.profile_key(inspection_type)
    fname = f"{project.replace(' ','_')}_{profile_name}_{date}_{suffix}.docx"
    out = REPORTS_DIR / fname
    doc.save(out)
    return out


def build_certificate(session):
    inspection_type = session.get("inspection_type", "")
    template = report_profiles.certificate_template_path(
        BASE_DIR, inspection_type
    )
    mode = session.get("certificate_mode", "none")
    if template is None or mode == "none":
        return None

    project = session.get("project_name", "—")
    address = session.get("address", "—")
    date = session.get("date", datetime.today().strftime("%Y-%m-%d"))
    groups = session.get("groups", [])
    profile_name = report_profiles.profile_key(inspection_type)
    exclusions = report_profiles.certificate_exclusions(groups)

    if profile_name == "anchor_annual":
        inspection_text = (
            "une inspection visuelle annuelle des systèmes d’ancrage, des "
            "lignes de vie, des bossoirs et des composantes accessibles"
        )
        validity = (
            f"Certificat valide à compter du {date} pour une période maximale "
            "de douze (12) mois, sous réserve des exclusions indiquées."
        )
    else:
        inspection_text = (
            "une inspection quinquennale des systèmes d’ancrage, des "
            "bossoirs et des composantes accessibles, incluant les essais "
            "applicables consignés au rapport"
        )
        validity = (
            f"Le présent certificat confirme la réalisation de l’inspection "
            f"quinquennale en date du {date}. Les inspections annuelles "
            "requises doivent continuer d’être effectuées."
        )

    body = (
        f"La présente attestation confirme que le système visé a fait l’objet "
        f"de {inspection_text}, conformément au rapport correspondant."
    )
    if mode == "with_exclusions":
        exclusion_text = "; ".join(exclusions) or (
            "les éléments expressément identifiés au rapport"
        )
        body += (
            " Sont exclus du présent certificat : "
            f"{exclusion_text}. Ils doivent demeurer hors service jusqu’aux "
            "correctifs et, lorsque requis, à une nouvelle inspection."
        )

    doc = Document(template)
    apply_replacements(
        doc,
        {
            "{{Project_Name}}": project,
            "{{Address_of_project}}": address,
            "{{Date}}": date,
            "{{Certificate_Body}}": body,
            "{{Certificate_Validity}}": validity,
        },
    )
    out = REPORTS_DIR / (
        f"{project.replace(' ', '_')}_Certificate_{profile_name}_{date}.docx"
    )
    doc.save(out)
    return out


# ══════════════════════════════════════════════════════════════════════════════
#  HANDLERS
# ══════════════════════════════════════════════════════════════════════════════


async def got_main_menu(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    choice = update.message.text.strip()
    if "Define" in choice or "project" in choice.lower():
        await update.message.reply_text(
            "➕ *Add new project*\n\nWhat is the *project name*?",
            parse_mode="Markdown", reply_markup=ReplyKeyboardRemove())
        return STATE_ADMIN_PROJECT_NAME
    else:
        chat_id = update.effective_chat.id
        save_session(chat_id, {"groups":[], "plans":[], "davit_detail":None,
                               "date": datetime.today().strftime("%Y-%m-%d")})
        projects = get_projects()
        if not projects:
            await update.message.reply_text("⚠️ No projects found. Use 'Define new project' first.")
            return STATE_MAIN_MENU
        buttons = [[p["name"]] for p in projects]
        await update.message.reply_text(
            "📋 Which project?",
            reply_markup=ReplyKeyboardMarkup(buttons, one_time_keyboard=True, resize_keyboard=True))
        return STATE_PROJECT_SELECT

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    clear_session(chat_id)
    await update.message.reply_text(
        "👷 *BSF Inspections – Report Bot*\n\nWhat would you like to do?",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(
            [["📁 Define new project"], ["📝 Write a report"]],
            one_time_keyboard=True, resize_keyboard=True))
    return STATE_MAIN_MENU

async def got_inspection_type(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    choice  = update.message.text.strip()
    session["inspection_type"] = choice
    save_session(chat_id, session)

    projects = get_projects()
    if not projects:
        await update.message.reply_text("⚠️ No projects in database. Use /addproject to add one.")
        return ConversationHandler.END

    buttons = [[p["name"]] for p in projects]
    await update.message.reply_text(
        f"✅ *{choice}*\n\nWhich project?",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(buttons, one_time_keyboard=True, resize_keyboard=True))
    return STATE_PROJECT_SELECT

async def got_project_select(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id  = update.effective_chat.id
    session  = load_session(chat_id)
    choice   = update.message.text.strip()
    projects = get_projects()

    project = next((p for p in projects if p["name"] == choice), None)
    if not project:
        await update.message.reply_text("❌ Project not found. Try again.")
        return STATE_PROJECT_SELECT

    # Load project data into session
    session["project_name"]  = project["name"]
    session["address"]       = project["address"]
    session["plans"]         = project.get("plans", [])
    session["davit_detail"]  = project.get("davit_detail")
    save_session(chat_id, session)

    plans_count = len(session["plans"])
    davit       = "✅" if session["davit_detail"] else "—"

    await update.message.reply_text(
        f"✅ *{project['name']}*\n"
        f"📍 {project['address']}\n"
        f"🗺 Plans: {plans_count} | Davit detail: {davit}\n\n"
        "📸 Send the *first inspection photo*.",
        parse_mode="Markdown", reply_markup=ReplyKeyboardRemove())
    return STATE_PHOTO

async def got_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id    = update.effective_chat.id
    session    = load_session(chat_id)
    photo_file = await update.message.photo[-1].get_file()
    total      = sum(len(g["photos"]) for g in session.get("groups",[]))
    photo_path = str(PHOTOS_DIR / f"{chat_id}_photo_{total+1}.jpg")
    await photo_file.download_to_drive(photo_path)
    ctx.user_data["pending_photo_path"] = photo_path

    # Show existing groups as quick options
    groups = session.get("groups", [])
    existing_types = [g.get("element_type","") for g in groups]

    if groups:
        # Build buttons with index prefix for reliable matching
        buttons = [["[" + str(i+1) + "] " + g.get("element_type","Group") + " (" + str(len(g["photos"])) + " photos)"] for i, g in enumerate(groups)]
        buttons.append(["🆕 New element type"])
        await update.message.reply_text(
            "📷 Photo received!\n\nAdd to which group?",
            reply_markup=ReplyKeyboardMarkup(buttons, one_time_keyboard=True, resize_keyboard=True))
        return STATE_GROUP_OR_ADD
    else:
        await update.message.reply_text(
            "📷 First photo!\n\n🔩 What *element type* is this?",
            parse_mode="Markdown",
            reply_markup=ReplyKeyboardMarkup(ELEMENT_TYPES, one_time_keyboard=True, resize_keyboard=True))
        return STATE_ELEMENT_TYPE

async def got_group_or_add(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    choice  = update.message.text.strip()
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    if "Add" in choice:
        session["groups"][-1]["photos"].append({"path": ctx.user_data["pending_photo_path"]})
        save_session(chat_id, session)
        n = len(session["groups"][-1]["photos"])
        await update.message.reply_text(
            f"✅ Photo added ({n} total).\n\n📸 Send another, /done to finish, or /remove to delete last photo.",
            parse_mode="Markdown", reply_markup=ReplyKeyboardRemove())
        return STATE_PHOTO
    else:
        await update.message.reply_text(
            "🆕 New group!\n\n🔩 What *element type* is this?",
            parse_mode="Markdown",
            reply_markup=ReplyKeyboardMarkup(ELEMENT_TYPES, one_time_keyboard=True, resize_keyboard=True))
        return STATE_ELEMENT_TYPE

ELEMENT_ID_QUESTIONS = {
    "Anchor":        "🔢 *Anchor number?*\n\n_e.g. AB1 · AB-05 · AN2…_",
    "Davit":         "🔢 *Davit number?*\n\n_e.g. D1 · D-03 · Bossoir #2…_",
    "Cable":         "🔢 *Cable / lifeline ID?*\n\n_e.g. C1 · Câble Nord…_",
    "Base / Socket": "🔢 *Base / socket number?*\n\n_e.g. S1 · Socle #3…_",
    "Facade":        "📍 *Facade zone?*\n\n_e.g. North · Level 5 · Grid B…_",
    "Roof":          "📍 *Roof zone?*\n\n_e.g. NE corner · Near stairwell…_",
    "Other":         "📍 *Describe the location:*\n\n_e.g. Mechanical room · Level 3…_",
}

async def got_element_type(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    element = update.message.text.strip()
    ctx.user_data["element_type"] = element
    ctx.user_data["location"] = element

    # Check if group with this element type already exists
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    for i, g in enumerate(session.get("groups", [])):
        if g.get("element_type","").strip() == element.strip():
            # Add photo directly to existing group
            # Store which group to add to
            ctx.user_data["add_to_group_idx"] = i
            await update.message.reply_text(
                "What is the status of this element?",
                reply_markup=ReplyKeyboardMarkup([
                    ["✅ Acceptable"],
                    ["🔧 Réparation requise"],
                    ["🔄 Remplacement requis"],
                    ["❌ Rejeté"],
                ], one_time_keyboard=True, resize_keyboard=True))
            return STATE_ELEMENT_STATUS

    # No existing group — ask status first
    ctx.user_data.pop("add_to_group_idx", None)
    await update.message.reply_text(
        "What is the status of this element?",
        reply_markup=ReplyKeyboardMarkup([
            ["✅ Acceptable"],
            ["🔧 Réparation requise"],
            ["🔄 Remplacement requis"],
            ["❌ Rejeté"],
        ], one_time_keyboard=True, resize_keyboard=True))
    return STATE_ELEMENT_STATUS

async def got_element_id(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    # Element ID removed - go straight to problem
    ctx.user_data["location"] = ctx.user_data.get("element_type","")
    await update.message.reply_text(
        "🔍 Describe the observation, or tap if no visible issue.",
        reply_markup=ReplyKeyboardMarkup([["⏭ No visible issue"]], one_time_keyboard=True, resize_keyboard=True))
    return STATE_PROBLEM

async def got_problem(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    problem = update.message.text.strip()
    if problem.lower() in ("skip","⏭ no visible issue","no visible issue"): problem = ""
    await update.message.reply_text("🤖 Analysing photo with AI…")
    try:
        img_bytes = Path(ctx.user_data["pending_photo_path"]).read_bytes()
        ai = analyse_photo(img_bytes, ctx.user_data.get("element_type","Unknown"),
                          ctx.user_data.get("location","Unknown"), problem)
    except Exception as e:
        log.error(f"API error: {e}")
        ai = {"caption_fr":"Observation à compléter","caption_en":"Observation to be completed","severity":"minor"}
    sev = SEVERITY_MAP.get(ai.get("severity","ok"),"")
    ctx.user_data["pending_ai"] = ai
    await update.message.reply_text(
        f"*Suggested caption (FR):* {ai.get('caption_fr')}\n*Severity:* {sev}\n\nAccept or type your own:",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup([[f"✅ {ai.get('caption_fr')}"],["✏️ Write my own"]], one_time_keyboard=True, resize_keyboard=True))
    return STATE_GROUP_CAPTION_FR


async def got_element_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    status  = update.message.text.strip()

    idx = ctx.user_data.get("add_to_group_idx")
    if idx is not None:
        # Adding to existing group
        session["groups"][idx]["photos"].append({
            "path":   ctx.user_data["pending_photo_path"],
            "status": status,
        })
        save_session(chat_id, session)
        n = len(session["groups"][idx]["photos"])
        await update.message.reply_text(
            "✅ Added (" + str(n) + " photos).\n\n📸 Send next photo or /done.",
            reply_markup=ReplyKeyboardRemove())
        ctx.user_data.pop("add_to_group_idx", None)
        return STATE_PHOTO
    else:
        # New group — status stored, then ask for caption
        ctx.user_data["pending_status"] = status
        await update.message.reply_text(
            "🤖 Analysing photo with AI…")
        try:
            img_bytes = Path(ctx.user_data["pending_photo_path"]).read_bytes()
            ai = analyse_photo(img_bytes, ctx.user_data.get("element_type","Unknown"),
                              ctx.user_data.get("location","Unknown"), "")
        except Exception as e:
            log.error(f"API error: {e}")
            ai = {"caption_fr":"Observation à compléter","caption_en":"Observation to be completed","severity":"minor"}
        ctx.user_data["pending_ai"] = ai

        # Get auto caption from inspection type + element type
        inspection_type = session.get("inspection_type", "")
        element_type    = ctx.user_data.get("element_type", "")
        auto_caption    = CAPTION_MAP.get((inspection_type, element_type), ai.get("caption_fr","Observation à compléter"))

        # Build caption options
        options = [
            ["✅ " + auto_caption],
            ["✅ " + ai.get("caption_fr","")] if ai.get("caption_fr","") != auto_caption else None,
            ["✏️ Write my own"],
        ]
        options = [o for o in options if o]  # remove None

        ctx.user_data["auto_caption"] = auto_caption
        await update.message.reply_text(
            "Choose caption or write your own:",
            reply_markup=ReplyKeyboardMarkup(options, one_time_keyboard=True, resize_keyboard=True))
        return STATE_GROUP_CAPTION_FR

async def got_group_caption_fr(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    ai   = ctx.user_data.get("pending_ai", {})
    if text == "✏️ Write my own":
        await update.message.reply_text("✏️ Type your French caption:", reply_markup=ReplyKeyboardRemove())
        return STATE_GROUP_CAPTION_FR
    elif text.startswith("✅ "):
        ctx.user_data["final_caption_fr"] = text[2:].strip()
    else:
        ctx.user_data["final_caption_fr"] = text
    # Skip EN caption — use auto-translation from FR
    caption_fr = ctx.user_data.get("final_caption_fr","")
    caption_en = ai.get("caption_en","")
    chat_id    = update.effective_chat.id
    session    = load_session(chat_id)
    session.setdefault("groups",[]).append({
        "element_type": ctx.user_data.get("element_type",""),
        "caption_fr": caption_fr, "caption_en": caption_en,
        "severity": ai.get("severity","ok"),
        "photos": [{"path": ctx.user_data["pending_photo_path"], "status": ctx.user_data.get("pending_status","✅ Acceptable")}],
    })
    save_session(chat_id, session)
    await update.message.reply_text(
        "✅ Group " + str(len(session['groups'])) + " created.\n\n📸 Send next photo or type */done*.",
        parse_mode="Markdown", reply_markup=ReplyKeyboardRemove())
    return STATE_PHOTO


async def cmd_remove_last(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    groups  = session.get("groups", [])
    if not groups:
        await update.message.reply_text("No photos to remove.")
        return
    last_group = groups[-1]
    photos     = last_group.get("photos", [])
    if len(photos) > 1:
        photos.pop()
        last_group["photos"] = photos
        save_session(chat_id, session)
        await update.message.reply_text(
            "Removed. Group now has " + str(len(photos)) + " photo(s).\n\nSend next photo or /done.",
            reply_markup=ReplyKeyboardRemove())
    elif len(photos) == 1:
        groups.pop()
        session["groups"] = groups
        save_session(chat_id, session)
        await update.message.reply_text(
            "Removed. " + str(len(groups)) + " group(s) remaining.\n\nSend next photo or /done.",
            reply_markup=ReplyKeyboardRemove())
    else:
        await update.message.reply_text("No photos to remove.")

async def cmd_done(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    groups  = session.get("groups",[])
    total   = sum(len(g["photos"]) for g in groups)
    if not groups:
        await update.message.reply_text("⚠️ No photos yet.")
        return STATE_PHOTO
    await update.message.reply_text(
        f"📝 Building report for *{session.get('project_name')}*…\n{len(groups)} group(s), {total} photo(s)",
        parse_mode="Markdown")
    try:
        report_fr = build_report(session, "fr")
        await update.message.reply_document(open(report_fr,"rb"), filename=report_fr.name, caption="🇫🇷 Rapport Word")
        try:
            pdf_out = report_fr.with_suffix(".pdf")
            docx_to_pdf(str(report_fr), str(pdf_out))
            if pdf_out.exists():
                await update.message.reply_document(open(pdf_out,"rb"), filename=pdf_out.name, caption="🇫🇷 Rapport PDF")
            else:
                await update.message.reply_text("PDF not created (file missing)")
        except Exception as e:
            log.error(f"PDF error: {e}")
            await update.message.reply_text(f"PDF error: {e}")
        await update.message.reply_text("✅ Rapport envoyé!\nType /start for a new inspection.")
    except Exception as e:
        log.error(f"Report error: {e}")
        await update.message.reply_text(f"❌ Error: {e}")
    clear_session(chat_id)
    return ConversationHandler.END


# ── Admin: Add project ─────────────────────────────────────────────────────
async def cmd_addproject(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "➕ *Add new project*\n\nWhat is the *project name*?",
        parse_mode="Markdown", reply_markup=ReplyKeyboardRemove())
    return STATE_ADMIN_PROJECT_NAME

async def admin_got_name(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["new_project"] = {"id": f"proj_{int(datetime.now().timestamp())}", "name": update.message.text.strip(), "plans": [], "davit_detail": None}
    await update.message.reply_text("📍 What is the *site address*?", parse_mode="Markdown")
    return STATE_ADMIN_PROJECT_ADDRESS

async def admin_got_address(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["new_project"]["address"] = update.message.text.strip()
    await update.message.reply_text(
        "🗺 Send *floor plan photos* for this project, or tap Skip.",
        reply_markup=ReplyKeyboardMarkup([["⏭ Skip plans"]], one_time_keyboard=True, resize_keyboard=True))
    return STATE_ADMIN_PROJECT_PLANS

async def admin_got_plan(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    photo_file = await update.message.photo[-1].get_file()
    proj_id    = ctx.user_data["new_project"]["id"]
    plan_idx   = len(ctx.user_data["new_project"]["plans"]) + 1
    plan_path  = str(PHOTOS_DIR / f"{proj_id}_plan_{plan_idx}.jpg")
    await photo_file.download_to_drive(plan_path)
    ctx.user_data["new_project"]["plans"].append({"path": plan_path, "ai": {
        "caption_fr": f"Plan {plan_idx}", "caption_en": f"Plan {plan_idx}",
        "severity": "ok", "detail_fr": "", "detail_en": ""}})
    await update.message.reply_text(
        f"✅ Plan {plan_idx} saved!",
        reply_markup=ReplyKeyboardMarkup([["📎 Add another plan","⏭ Skip plans"]], one_time_keyboard=True, resize_keyboard=True))
    return STATE_ADMIN_PROJECT_PLANS

async def admin_skip_plans(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🔩 Send the *davit detail photo*, or tap Skip.",
        reply_markup=ReplyKeyboardMarkup([["⏭ Skip davit"]], one_time_keyboard=True, resize_keyboard=True))
    return STATE_ADMIN_PROJECT_DAVIT

async def admin_got_davit(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    photo_file = await update.message.photo[-1].get_file()
    proj_id    = ctx.user_data["new_project"]["id"]
    davit_path = str(PHOTOS_DIR / f"{proj_id}_davit.jpg")
    await photo_file.download_to_drive(davit_path)
    ctx.user_data["new_project"]["davit_detail"] = davit_path
    return await admin_save_project(update, ctx)

async def admin_skip_davit(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    return await admin_save_project(update, ctx)

async def admin_save_project(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    db = load_db()
    db["projects"].append(ctx.user_data["new_project"])
    save_db(db)
    name = ctx.user_data["new_project"]["name"]
    msg = "✅ Project *" + name + "* saved!\n\nWhat would you like to do next?"
    await update.message.reply_text(
        msg, parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(
            [["📁 Define new project"], ["📝 Write a report"]],
            one_time_keyboard=True, resize_keyboard=True))
    return STATE_MAIN_MENU

async def cmd_projects(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    projects = get_projects()
    if not projects:
        await update.message.reply_text("No projects yet. Use /addproject.")
        return
    msg = "📋 *Projects:*\n\n"
    for i, p in enumerate(projects, 1):
        plans = len(p.get("plans",[]))
        davit = "✅" if p.get("davit_detail") else "—"
        msg  += f"{i}. *{p['name']}*\n   📍 {p['address']}\n   Plans: {plans} | Davit: {davit}\n\n"
    await update.message.reply_text(msg, parse_mode="Markdown")

async def cmd_cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    clear_session(update.effective_chat.id)
    await update.message.reply_text("❌ Cancelled. Type /start to begin.", reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END

async def cmd_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    session = load_session(chat_id)
    if not session:
        await update.message.reply_text("No active inspection. Type /start.")
        return
    groups = session.get("groups",[])
    total  = sum(len(g["photos"]) for g in groups)
    msg    = (f"📊 *Status*\nProject: {session.get('project_name','—')}\n"
              f"Type: {session.get('inspection_type','—')}\n"
              f"Groups: {len(groups)} | Photos: {total}\n\n")
    for i, g in enumerate(groups, 1):
        msg += f"  {i}. {g.get('caption_en','—')} ({len(g['photos'])} photos)\n"
    await update.message.reply_text(msg, parse_mode="Markdown")


# ══════════════════════════════════════════════════════════════════════════════
def main():
    app  = Application.builder().token(TELEGRAM_TOKEN).build()

    # Main inspection flow
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", cmd_start)],
        states={
            STATE_MAIN_MENU:        [MessageHandler(filters.TEXT & ~filters.COMMAND, got_main_menu)],
            STATE_ADMIN_PROJECT_NAME:    [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_got_name)],
            STATE_ADMIN_PROJECT_ADDRESS: [MessageHandler(filters.TEXT & ~filters.COMMAND, admin_got_address)],
            STATE_ADMIN_PROJECT_PLANS:   [MessageHandler(filters.PHOTO, admin_got_plan),
                                          MessageHandler(filters.TEXT & ~filters.COMMAND, admin_skip_plans)],
            STATE_ADMIN_PROJECT_DAVIT:   [MessageHandler(filters.PHOTO, admin_got_davit),
                                          MessageHandler(filters.TEXT & ~filters.COMMAND, admin_skip_davit)],
            STATE_INSPECTION_TYPE:  [MessageHandler(filters.TEXT & ~filters.COMMAND, got_inspection_type)],
            STATE_PROJECT_SELECT:   [MessageHandler(filters.TEXT & ~filters.COMMAND, got_project_select)],
            STATE_PHOTO:            [MessageHandler(filters.PHOTO, got_photo), CommandHandler("done", cmd_done)],
            STATE_GROUP_OR_ADD:     [MessageHandler(filters.TEXT & ~filters.COMMAND, got_group_or_add)],
            STATE_ELEMENT_TYPE:     [MessageHandler(filters.TEXT & ~filters.COMMAND, got_element_type)],
            STATE_ELEMENT_ID:       [MessageHandler(filters.TEXT & ~filters.COMMAND, got_element_id)],
            STATE_ELEMENT_STATUS:   [MessageHandler(filters.TEXT & ~filters.COMMAND, got_element_status)],
            STATE_PROBLEM:          [MessageHandler(filters.TEXT & ~filters.COMMAND, got_problem)],
            STATE_GROUP_CAPTION_FR: [MessageHandler(filters.TEXT & ~filters.COMMAND, got_group_caption_fr)],

        },
        fallbacks=[CommandHandler("cancel", cmd_cancel), CommandHandler("done", cmd_done), CommandHandler("remove", cmd_remove_last)],
        allow_reentry=True,
    )

    app.add_handler(conv)
    app.add_handler(CommandHandler("projects", cmd_projects))
    app.add_handler(CommandHandler("status", cmd_status))
    app.add_handler(CommandHandler("remove", cmd_remove_last))

    log.info("🚀 BSF Inspection Bot running…")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
