"""Fallback certificate generator used when a DOCX template cannot be opened."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import inspection_bot as bot

log = logging.getLogger(__name__)


def _safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return value.strip("._") or "inspection"


def _build_simple_certificate(session: dict) -> Path | None:
    inspection_type = session.get("inspection_type", "")
    mode = session.get("certificate_mode", "none")
    if mode == "none" or not bot.report_profiles.is_anchor(inspection_type):
        return None

    project = session.get("project_name", "—")
    address = session.get("address", "—")
    date = session.get("date", "—")
    profile_name = bot.report_profiles.profile_key(inspection_type)
    exclusions = bot.report_profiles.certificate_exclusions(session.get("groups", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CERTIFICAT D’INSPECTION")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Système d’accès suspendu")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Projet", project),
        ("Adresse", address),
        ("Date de l’inspection", date),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = str(value)

    if profile_name == "anchor_annual":
        body = (
            "La présente attestation confirme qu’une inspection visuelle annuelle "
            "des composantes accessibles du système d’accès suspendu a été réalisée "
            "à la date indiquée ci-dessus."
        )
        validity = "Validité maximale de douze (12) mois, sous réserve des exclusions indiquées."
    else:
        body = (
            "La présente attestation confirme qu’une inspection quinquennale des "
            "composantes accessibles du système d’accès suspendu a été réalisée, "
            "incluant les essais applicables consignés au rapport."
        )
        validity = (
            "Les inspections annuelles requises doivent continuer d’être effectuées."
        )

    doc.add_paragraph(body)

    if mode == "with_exclusions":
        exclusion_text = "; ".join(exclusions) or "les éléments identifiés au rapport"
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Exclusions : ")
        run.bold = True
        paragraph.add_run(
            exclusion_text
            + ". Ces éléments doivent demeurer hors service jusqu’aux correctifs "
            "et, lorsque requis, à une nouvelle inspection."
        )
    else:
        doc.add_paragraph("Aucune exclusion n’a été enregistrée dans le rapport.")

    doc.add_paragraph(validity)
    doc.add_paragraph()
    doc.add_paragraph("Préparé par : ______________________________")
    doc.add_paragraph("Signature : _________________________________")
    doc.add_paragraph("Date : _____________________________________")

    bot.REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output = bot.REPORTS_DIR / (
        f"{_safe_name(project)}_Certificate_{profile_name}_{_safe_name(str(date))}.docx"
    )
    doc.save(output)
    return output


def install_certificate_fallback() -> None:
    original = bot.build_certificate

    def build_certificate(session: dict):
        try:
            return original(session)
        except Exception:
            log.exception("Template certificate generation failed; using fallback")
            return _build_simple_certificate(session)

    bot.build_certificate = build_certificate
